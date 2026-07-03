"""
report_docx.py — конвертация markdown-отчёта в .docx (Word) через python-docx.

Без pandoc и без LLM. Понимает заголовки (#/##/###), списки (- / •), жирный **...**,
курсив _..._ и обычные абзацы. Используется комбо-ботом (report.docx) и может
применяться к отчётам пульсового бота.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


def _add_runs(paragraph, text: str):
    """Разбирает **жирный** и _курсив_ внутри строки."""
    for part in re.split(r"(\*\*[^*]+\*\*|_[^_]+_)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("_") and part.endswith("_"):
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def md_to_docx(md_text: str, out_path: str | Path, title: str | None = None) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=0)

    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.lstrip().startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line.lstrip()[2:])
        elif re.match(r"^_.*_$", line.strip()):
            p = doc.add_paragraph()
            r = p.add_run(line.strip().strip("_")); r.italic = True
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    out_path = Path(out_path)
    doc.save(str(out_path))
    return str(out_path)


def md_file_to_docx(md_path: str | Path, out_path: str | Path | None = None) -> str:
    md_path = Path(md_path)
    out_path = out_path or md_path.with_suffix(".docx")
    return md_to_docx(md_path.read_text(encoding="utf-8"), out_path)


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", required=True)
    ap.add_argument("--out", default=None)
    a = ap.parse_args()
    print(md_file_to_docx(a.md, a.out))
