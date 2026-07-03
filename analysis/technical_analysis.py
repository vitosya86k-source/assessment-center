"""
Технический анализ транскриптов Assessment Center
Многоуровневая предобработка текстов для извлечения психологических паттернов

Архитектура:
1. Чтение Word файлов
2. Анализ дефлюенси (НА СЫРОМ ТЕКСТЕ)
3. Очистка текста
4. Извлечение критических инцидентов
5. Анализ по 12 модулям
6. Кросс-валидация через упражнения
7. Создание компактной JSON выжимки (800-1000 слов)
"""

import re
import os
import json
import logging
from typing import List, Dict, Any, Tuple
from collections import Counter, defaultdict
from functools import lru_cache
import inspect
from collections import namedtuple

# Патч для pymorphy2
if not hasattr(inspect, 'getargspec'):
    ArgSpec = namedtuple('ArgSpec', ['args', 'varargs', 'keywords', 'defaults'])
    def getargspec(func):
        spec = inspect.getfullargspec(func)
        return ArgSpec(args=spec.args, varargs=spec.varargs, keywords=spec.varkw, defaults=spec.defaults)
    inspect.getargspec = getargspec

try:
    import pymorphy2
    morph = pymorphy2.MorphAnalyzer()
    HAS_PYMORPHY = True
except Exception as e:
    HAS_PYMORPHY = False
    morph = None
    logging.warning(f"pymorphy2 недоступен: {e}")

# Sentence-BERT для семантического анализа
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    semantic_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
    HAS_SEMANTIC = True
    print("✅ Sentence-BERT модель загружена для семантического анализа")
except Exception as e:
    HAS_SEMANTIC = False
    semantic_model = None
    print(f"⚠️ Sentence-transformers недоступен: {e}. Будет использован базовый анализ")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("python-docx не установлен")

# Импорт словарей маркеров
from markers_dictionaries import (
    MBTI_MARKERS, PAVLOV_MARKERS, OCEAN_MARKERS, RADICALS_MARKERS,
    DARK_TETRAD_MARKERS, POTENTIAL_MARKERS, STRESS_MARKERS,
    PRISM_MARKERS, SHADOW_LOGIC, DISFLUENCY_CATEGORIES,
    COGNITIVE_FUNCTIONS_PATTERNS, STOP_WORDS, MIN_FREQUENCIES
)

# Настройка логирования
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


# ============================================================================
# ЧТЕНИЕ WORD ФАЙЛОВ
# ============================================================================

def read_docx(file_path: str) -> str:
    """
    Читает текст из Word документа
    
    Args:
        file_path: путь к .docx файлу
        
    Returns:
        извлеченный текст
    """
    if not HAS_DOCX:
        raise ImportError("python-docx не установлен. Установите: pip install python-docx")
    
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)
        logger.info(f"Прочитано {len(paragraphs)} параграфов, {len(text.split())} слов")
        return text
    except Exception as e:
        logger.error(f"Ошибка чтения {file_path}: {e}")
        raise

def extract_participant_text_from_docx(file_path: str) -> str:
    """
    Извлекает ТОЛЬКО текст участника из секции '📝 Текст участника (для анализа)'
    
    Args:
        file_path: путь к .docx файлу
        
    Returns:
        текст участника или пустая строка
    """
    if not HAS_DOCX:
        return ""
    
    try:
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs]
        
        # Ищем заголовок "Текст участника"
        participant_text = []
        in_participant_section = False
        
        for para in paragraphs:
            # Начало секции участника
            if "Текст участника" in para or "текст участника" in para.lower():
                in_participant_section = True
                continue
            
            # Конец секции (начало следующей секции или полный транскрипт)
            if in_participant_section and ("Полный транскрипт" in para or "ТЕКСТЫ ВСЕХ СПИКЕРОВ" in para):
                break
            
            # Собираем текст
            if in_participant_section and para.strip() and len(para.strip()) > 20:
                # Пропускаем заголовки и метаданные
                if not para.startswith("📊") and not para.startswith("⚡") and not para.startswith("💬"):
                    participant_text.append(para.strip())
        
        result = "\n".join(participant_text)
        if result:
            logger.info(f"✅ Извлечен текст участника: {len(result.split())} слов из {os.path.basename(file_path)}")
        return result
        
    except Exception as e:
        logger.error(f"Ошибка извлечения текста участника из {file_path}: {e}")
        return ""


def extract_speech_metrics_from_docx(file_path: str) -> Dict[str, Any]:
    """
    Извлекает метрики речи из Word документа (если есть)
    
    Args:
        file_path: путь к .docx файлу
        
    Returns:
        словарь с метриками речи или None
    """
    if not HAS_DOCX:
        return None
    
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        metrics = {}
        
        # Ищем скорость речи
        wpm_match = re.search(r'⚡\s*Скорость речи[:\s]+(\d+[\.,]?\d*)\s*слов/мин', text)
        if wpm_match:
            metrics['words_per_minute'] = float(wpm_match.group(1).replace(',', '.'))
        
        # Ищем среднюю длину реплики
        avg_match = re.search(r'📝\s*Средняя длина реплики[:\s]+(\d+[\.,]?\d*)\s*слов', text)
        if avg_match:
            metrics['avg_words_per_utterance'] = float(avg_match.group(1).replace(',', '.'))
        
        # Ищем долю времени речи
        time_match = re.search(r'⏱️\s*Доля времени речи[:\s]+(\d+[\.,]?\d*)%', text)
        if time_match:
            metrics['speaking_time_percentage'] = float(time_match.group(1).replace(',', '.'))
        
        # Ищем общее количество реплик
        utterances_match = re.search(r'💬\s*Реплик всего[:\s]+(\d+)', text)
        if utterances_match:
            metrics['total_utterances'] = int(utterances_match.group(1))
        
        if metrics:
            logger.info(f"✅ Извлечены метрики речи из {os.path.basename(file_path)}: WPM={metrics.get('words_per_minute', 'N/A')}")
            return metrics
        else:
            return None
        
    except Exception as e:
        logger.error(f"Ошибка извлечения метрик из {file_path}: {e}")
        return None


def read_multiple_files(file_paths: List[str]) -> Tuple[List[Dict[str, str]], List[Dict[str, Any]], List[str]]:
    """
    Читает несколько файлов упражнений, извлекает метрики речи И тексты участников
    
    Args:
        file_paths: список путей к файлам
        
    Returns:
        кортеж: (список словарей с текстами, список метрик речи, список текстов участников)
    """
    exercises = []
    all_metrics = []
    participant_texts = []
    
    for idx, path in enumerate(file_paths, 1):
        try:
            text = read_docx(path)
            exercises.append({
                "exercise_id": idx,
                "file_name": os.path.basename(path),
                "text": text,
                "word_count": len(text.split())
            })
            
            # Извлекаем метрики речи (если есть)
            metrics = extract_speech_metrics_from_docx(path)
            if metrics:
                metrics['exercise_id'] = idx
                metrics['file_name'] = os.path.basename(path)
                all_metrics.append(metrics)
            
            # Извлекаем ТОЛЬКО текст участника (если есть секция)
            participant_text = extract_participant_text_from_docx(path)
            if participant_text:
                participant_texts.append(participant_text)
            
        except Exception as e:
            logger.error(f"Не удалось прочитать {path}: {e}")
            continue
    
    return exercises, all_metrics, participant_texts


# ============================================================================
# АНАЛИЗ ДЕФЛЮЕНСИ (НА СЫРОМ ТЕКСТЕ!)
# ============================================================================

def analyze_disfluencies(text: str) -> Dict[str, Any]:
    """
    Анализ речевых дефлюенси: паузы, филлеры, повторы
    КРИТИЧНО: вызывается ДО очистки текста!
    
    Args:
        text: сырой текст транскрипта
        
    Returns:
        профиль дефлюенси
    """
    words = text.lower().split()
    total_words = len(words)
    
    if total_words < 50:
        return {"baseline_filler_rate": 0, "note": "Текст слишком короткий"}
    
    # 1. Подсчет филлеров по категориям
    filler_counts = defaultdict(int)
    filler_positions = defaultdict(list)
    
    for category, data in DISFLUENCY_CATEGORIES.items():
        markers = data["markers"]
        for i, word in enumerate(words):
            if any(marker in word for marker in markers):
                filler_counts[category] += 1
                filler_positions[category].append(i)
    
    total_fillers = sum(filler_counts.values())
    baseline_filler_rate = total_fillers / total_words if total_words > 0 else 0
    
    # 2. Распределение типов филлеров
    filler_distribution = {}
    if total_fillers > 0:
        for category, count in filler_counts.items():
            filler_distribution[category] = round(count / total_fillers, 2)
    
    # 3. Поиск пауз (если есть метки)
    pause_count = len(re.findall(r'\[пауза\]|\[длинная пауза\]|\.\.\.', text))
    
    # 4. Поиск повторов (простая эвристика)
    repeat_count = 0
    for i in range(len(words) - 1):
        if words[i] == words[i+1] and len(words[i]) > 2:
            repeat_count += 1
    
    # 5. Самокоррекция (паттерны типа "я... то есть мы")
    self_corrections = len(re.findall(r'\.\.\.\s*(то есть|нет|вернее|точнее)', text.lower()))
    
    return {
        "baseline_filler_rate": round(baseline_filler_rate, 3),
        "baseline_interpretation": _interpret_filler_rate(baseline_filler_rate),
        "total_fillers": total_fillers,
        "filler_types_distribution": filler_distribution,
        "filler_counts_by_category": dict(filler_counts),
        "pause_count": pause_count,
        "repeat_count": repeat_count,
        "self_correction_count": self_corrections
    }


def _interpret_filler_rate(rate: float) -> str:
    """Интерпретация уровня дефлюенси"""
    if rate < 0.05:
        return "Very low disfluency - possibly rehearsed or very confident"
    elif rate < 0.10:
        return "Normal spontaneous speech"
    elif rate < 0.15:
        return "Moderate disfluency - thinking process visible"
    else:
        return "High disfluency - significant cognitive load or discomfort"


def analyze_contextual_disfluency(text: str, context_keywords: List[str]) -> float:
    """
    Анализ дефлюенси в конкретном контексте (для определения зон дискомфорта)
    
    Args:
        text: текст
        context_keywords: ключевые слова контекста (напр. "конфликт", "делегирование")
        
    Returns:
        уровень дефлюенси в этом контексте
    """
    sentences = re.split(r'[.!?]+', text)
    context_sentences = []
    
    for sent in sentences:
        if any(keyword in sent.lower() for keyword in context_keywords):
            context_sentences.append(sent)
    
    if not context_sentences:
        return 0.0
    
    context_text = " ".join(context_sentences)
    result = analyze_disfluencies(context_text)
    
    return result.get("baseline_filler_rate", 0.0)


# ============================================================================
# ОЧИСТКА ТЕКСТА (ПОСЛЕ АНАЛИЗА ДЕФЛЮЕНСИ!)
# ============================================================================

def extract_assessor_discussion(text: str) -> tuple[str, dict]:
    """
    Извлекает обсуждение ассессоров из конца файла (если есть)
    
    Обсуждение обычно:
    - Идет В КОНЦЕ после упражнений (последние 20-30% текста)
    - Упоминаются КОМПЕТЕНЦИИ, ОЦЕНКИ, БАЛЛЫ
    - НЕ путать с ролевыми упражнениями (там тоже 2 спикера, но с начала)
    
    Returns:
        (text_without_discussion, discussion_data)
    """
    words = text.split()
    total_words = len(words)
    
    if total_words < 500:  # Слишком короткий текст
        return text, None
    
    # Проверяем ПОСЛЕДНЮЮ ТРЕТЬ текста на маркеры обсуждения
    last_third_start = int(total_words * 0.67)  # Последние 33%
    last_third_text = ' '.join(words[last_third_start:])
    
    # Маркеры обсуждения компетенций
    discussion_markers = {
        'компетенция': 0,
        'компетенции': 0,
        'оценка': 0,
        'баллов': 0,
        'балла': 0,
        'проявил': 0,
        'проявила': 0,
        'проявляет': 0,
        'сильные стороны': 0,
        'зоны развития': 0,
        'поставим': 0,
        'поставить': 0,
        'ставлю': 0
    }
    
    for marker in discussion_markers:
        discussion_markers[marker] = last_third_text.lower().count(marker)
    
    total_markers = sum(discussion_markers.values())
    
    # Если мало маркеров - это НЕ обсуждение
    if total_markers < 3:
        return text, None
    
    # Ищем метки спикеров в последней трети
    speaker_pattern = r'\[[\d:\.]+\]\s*—\s*Спикер\s+(\d+)'
    matches_in_last = list(re.finditer(speaker_pattern, last_third_text))
    
    if len(matches_in_last) < 5:  # Мало реплик
        return text, None
    
    # Проверяем что это именно 2 спикера (ассессоры)
    speakers_in_last = set(m.group(1) for m in matches_in_last)
    
    if len(speakers_in_last) != 2:
        # Если не ровно 2 спикера - возможно это не обсуждение
        logger.info(f"⚠️ В последней трети {len(speakers_in_last)} спикеров - не похоже на обсуждение")
        return text, None
    
    # Ищем точку перехода к обсуждению
    # Идем от 50% до 80% текста и ищем где начинается концентрация маркеров
    best_split = None
    max_density = 0
    
    for split_pct in range(50, 85, 5):
        split_pos = int(total_words * split_pct / 100)
        candidate_discussion = ' '.join(words[split_pos:])
        
        # Плотность маркеров обсуждения
        markers_in_candidate = sum(candidate_discussion.lower().count(m) for m in discussion_markers)
        density = markers_in_candidate / len(candidate_discussion.split())
        
        if density > max_density:
            max_density = density
            best_split = split_pos
    
    if best_split is None or max_density < 0.005:  # Порог плотности
        return text, None
    
    # Разделяем текст
    interview_words = words[:best_split]
    discussion_words = words[best_split:]
    
    interview_text = ' '.join(interview_words)
    discussion_text = ' '.join(discussion_words)
    
    logger.info(f"📋 Найдено обсуждение ассессоров: {len(discussion_words)} слов ({int(len(discussion_words)/total_words*100)}% текста)")
    logger.info(f"   Маркеров обсуждения: {total_markers}")
    logger.info(f"   Спикеров в обсуждении: {len(speakers_in_last)}")
    
    # Извлекаем упомянутые компетенции
    competencies_mentioned = []
    comp_keywords = {
        'целостность': 'Целостность личности',
        'управление изменениями': 'Управление изменениями',
        'анализ': 'Анализ и принятие решений',
        'ориентация на результат': 'Ориентация на результат',
        'эффективное общение': 'Эффективное общение',
        'сотрудничество': 'Сотрудничество',
        'делегирование': 'Делегирование',
        'мотивация': 'Мотивация и развитие команды'
    }
    
    for keyword, full_name in comp_keywords.items():
        if keyword in discussion_text.lower():
            competencies_mentioned.append(full_name)
    
    discussion_data = {
        'text': discussion_text,
        'word_count': len(discussion_words),
        'marker_count': total_markers,
        'marker_density': max_density,
        'competencies_mentioned': competencies_mentioned,
        'split_position': best_split
    }
    
    return interview_text, discussion_data


def remove_exercise_legend(text: str) -> str:
    """
    Удаляет легенду упражнения (стандартные инструкции и описания)
    
    Легенда - это повторяющийся текст с описанием ситуации, который:
    - Идентичен для всех участников
    - Содержит инструкции, описание роли, контекст
    - НЕ является речью участника
    
    Args:
        text: исходный текст
        
    Returns:
        текст без легенды
    """
    # Маркеры начала диалога (после легенды)
    dialog_start_markers = [
        r'\[[\d:\.]+\]\s*—\s*Спикер',  # временная метка со спикером
        r'Как\s+прошла\s+встреча\?',
        r'Хотела\s+бы\s+поговорить',
        r'Присаживайтесь',
        r'Давайте\s+по\s+ним\s+пробежимся',
    ]
    
    # Ищем начало диалога
    dialog_start_pos = None
    for marker in dialog_start_markers:
        match = re.search(marker, text, re.IGNORECASE)
        if match:
            dialog_start_pos = match.start()
            logger.info(f"📍 Найдено начало диалога по маркеру на позиции {dialog_start_pos}")
            break
    
    # Если нашли начало диалога - отрезаем все до него
    if dialog_start_pos and dialog_start_pos > 100:  # минимум 100 символов легенды
        removed_text = text[:dialog_start_pos]
        text = text[dialog_start_pos:]
        logger.info(f"✂️ Удалена легенда: {len(removed_text)} символов ({len(removed_text.split())} слов)")
    
    return text.strip()


def clean_text(text: str) -> str:
    """
    Очистка текста от мусора
    Вызывается ПОСЛЕ analyze_disfluencies!
    
    Args:
        text: сырой текст
        
    Returns:
        очищенный текст
    """
    # Удаляем технические метки
    text = re.sub(r'\[пауза\]|\[длинная пауза\]|\[смех\]|\[кашель\]', '', text)
    text = re.sub(r'\d{1,2}:\d{2}:\d{2}', '', text)  # timestamps
    
    # Удаляем избыточные паузы
    text = re.sub(r'\.{3,}', '.', text)
    
    # Нормализуем пробелы
    text = re.sub(r'\s+', ' ', text)
    
    # Склеиваем разорванные предложения (простая эвристика)
    text = re.sub(r'\s+([,.:;!?])', r'\1', text)
    
    return text.strip()


def identify_speaker_patterns(text: str) -> Dict[str, int]:
    """
    Анализирует паттерны речи для определения кто есть кто
    Возвращает статистику по потенциальным спикерам
    """
    sentences = re.split(r'[.!?]+', text)
    speaker_stats = {
        'assessor': 0,
        'participant': 0,
        'role_player': 0,
        'unclear': 0
    }
    
    for sent in sentences:
        sent = sent.strip()
        if len(sent) < 10:
            continue
        
        sent_lower = sent.lower()
        
        # Явные маркеры ассессора
        assessor_markers = [
            '?', 'участник', 'кандидат', 'вы ', 'вам ', 'расскажите',
            'опишите', 'объясните', 'покажите', 'попробуйте',
            'показывает', 'демонстрирует', 'проявляет', 'наблюдаю'
        ]
        
        # Явные маркеры участника
        participant_markers = [
            ' я ', 'я думаю', 'я считаю', 'я чувствую', 'у меня',
            'мне кажется', 'в моей', 'мой ', 'мое ', 'делаю',
            'работаю', 'учусь', 'делал', 'работал'
        ]
        
        # Маркеры ролевой игры
        role_markers = [
            'директор:', 'менеджер:', 'клиент:', 'коллега:',
            'в роли', 'играю роль', 'я как директор', 'я как менеджер'
        ]
        
        if any(marker in sent_lower for marker in assessor_markers):
            speaker_stats['assessor'] += 1
        elif any(marker in sent_lower for marker in participant_markers):
            speaker_stats['participant'] += 1
        elif any(marker in sent_lower for marker in role_markers):
            speaker_stats['role_player'] += 1
        else:
            speaker_stats['unclear'] += 1
    
    return speaker_stats


def extract_speaker_turns(text: str) -> Dict[str, List[str]]:
    """
    Извлекает реплики каждого спикера из транскрипта с метками
    
    Поддерживаемые форматы:
    - [00:00:02.035] — Спикер 2
    - Спикер 0:
    - Speaker 1:
    
    Returns:
        словарь {speaker_id: [список_реплик]}
    """
    speakers = defaultdict(list)
    
    # Паттерн для разных форматов меток спикеров
    patterns = [
        r'\[[\d:\.]+\]\s*—\s*Спикер\s+(\d+)\s*\n+(.+?)(?=\[[\d:\.]+\]\s*—\s*Спикер\s+\d+|$)',  # [timestamp] — Спикер N
        r'Спикер\s+(\d+):\s*\n+(.+?)(?=Спикер\s+\d+:|$)',  # Спикер N:
        r'Speaker\s+(\d+):\s*\n+(.+?)(?=Speaker\s+\d+:|$)',  # Speaker N:
    ]
    
    found_speakers = False
    for pattern in patterns:
        matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
        if matches:
            found_speakers = True
            for speaker_id, content in matches:
                # Очищаем контент от лишних переносов
                content = re.sub(r'\s+', ' ', content.strip())
                speakers[f"Спикер {speaker_id}"].append(content)
            break
    
    if not found_speakers:
        logger.warning("⚠️ Метки спикеров не найдены в транскрипте")
    else:
        logger.info(f"✅ Найдено спикеров: {len(speakers)}")
        for speaker, turns in speakers.items():
            logger.info(f"  {speaker}: {len(turns)} реплик, {sum(len(t.split()) for t in turns)} слов")
    
    return dict(speakers)


def identify_participant_speaker(speakers_dict: Dict[str, List[str]]) -> str:
    """
    Определяет какой спикер является участником
    
    ПРОСТАЯ И НАДЕЖНАЯ ЛОГИКА:
    1. Исключаем спикера с большинством вопросов (это ассессор)
    2. Из оставшихся выбираем кто больше говорит (это участник)
    
    Args:
        speakers_dict: словарь {speaker_id: [реплики]}
        
    Returns:
        ID спикера-участника
    """
    if not speakers_dict:
        return None
    
    if len(speakers_dict) == 1:
        return list(speakers_dict.keys())[0]
    
    # Анализируем каждого спикера
    speaker_analysis = {}
    
    for speaker_id, turns in speakers_dict.items():
        full_text = ' '.join(turns)
        word_count = len(full_text.split())
        
        # Считаем ВОПРОСЫ - главный маркер ассессора
        question_count = full_text.count('?')
        
        # Считаем обращения "вы" - второй маркер ассессора
        vy_count = len(re.findall(r'\bвы\s+', full_text, re.IGNORECASE))
        
        # Маркеры ассессора
        assessor_phrases = len(re.findall(
            r'\bрасскажите\b|\bопишите\b|\bобъясните\b|\bскажите\b|\bприсаживайтесь\b',
            full_text,
            re.IGNORECASE
        ))
        
        # Итоговый "рейтинг ассессора"
        assessor_rating = question_count + (vy_count * 0.5) + (assessor_phrases * 2)
        
        speaker_analysis[speaker_id] = {
            'words': word_count,
            'turns': len(turns),
            'questions': question_count,
            'vy_count': vy_count,
            'assessor_phrases': assessor_phrases,
            'assessor_rating': assessor_rating
        }
        
        logger.info(
            f"  {speaker_id}: {word_count}w, "
            f"q={question_count}, вы={vy_count}, phrases={assessor_phrases}, "
            f"assessor_rating={assessor_rating:.1f}"
        )
    
    # ШАГ 1: Находим явного ассессора (больше всех вопросов)
    assessor_candidate = max(speaker_analysis.items(), key=lambda x: x[1]['assessor_rating'])
    
    # Если у кого-то явно много признаков ассессора - исключаем его
    if assessor_candidate[1]['assessor_rating'] > 5:
        logger.info(f"❌ Исключен ассессор: {assessor_candidate[0]} (rating={assessor_candidate[1]['assessor_rating']:.1f})")
        # Удаляем из рассмотрения
        remaining = {k: v for k, v in speaker_analysis.items() if k != assessor_candidate[0]}
    else:
        remaining = speaker_analysis
    
    # ШАГ 2: Из оставшихся выбираем кто больше говорит
    if remaining:
        participant = max(remaining.items(), key=lambda x: x[1]['words'])
        logger.info(f"✅ Участник определен: {participant[0]} ({participant[1]['words']} слов)")
        return participant[0]
    
    # Запасной вариант: просто берем кто больше говорит
    best_by_volume = max(speaker_analysis.items(), key=lambda x: x[1]['words'])
    logger.info(f"📊 Выбран по объему: {best_by_volume[0]} ({best_by_volume[1]['words']} слов)")
    return best_by_volume[0]


def extract_participant_speech(text: str) -> str:
    """
    УМНАЯ фильтрация речи участника от речи ассессора
    
    Поддерживает два режима:
    1. Если есть метки спикеров (Спикер N) - извлекает речь по ID
    2. Если меток нет - использует паттерн-анализ
    
    Args:
        text: полный транскрипт
        
    Returns:
        только речь участника
    """
    # Режим 1: Пытаемся извлечь по меткам спикеров
    speakers_dict = extract_speaker_turns(text)
    
    if speakers_dict:
        # Определяем кто из спикеров - участник
        participant_id = identify_participant_speaker(speakers_dict)
        
        if participant_id and participant_id in speakers_dict:
            participant_text = ' '.join(speakers_dict[participant_id])
            logger.info(f"📝 Извлечено речи участника ({participant_id}): {len(participant_text.split())} слов")
            return participant_text
    
    # Режим 2: Паттерн-анализ (если метки не найдены или не удалось определить)
    logger.info("🔍 Используем паттерн-анализ для фильтрации")
    
    # Сначала анализируем паттерны речи в тексте
    speaker_stats = identify_speaker_patterns(text)
    logger.info(f"Анализ паттернов: {speaker_stats}")
    
    sentences = re.split(r'[.!?]+', text)
    participant_sentences = []
    
    for sent in sentences:
        sent = sent.strip()
        if len(sent) < 15:
            continue
        
        sent_lower = sent.lower()
        
        # ===== ШАГ 1: СИЛЬНЫЕ МАРКЕРЫ УЧАСТНИКА (приоритет) =====
        # Если есть явные маркеры участника - берем предложение НЕСМОТРЯ на остальное
        strong_participant_markers = [
            ' я ', 'я думаю', 'я считаю', 'я чувствую',
            'я обычно', 'я стараюсь', 'я предпочитаю',
            'у меня', 'мне кажется', 'мне нравится',
            'в моей', 'в моем', 'моя ', 'мой ', 'мое ',
            'делаю', 'делал', 'делала', 'сделал', 'сделала',
            'работаю', 'работал', 'работала',
            'учусь', 'учился', 'училась', 'изучаю', 'изучал'
        ]
        
        has_strong_participant = any(pattern in sent_lower for pattern in strong_participant_markers)
        
        # ===== ШАГ 2: СИЛЬНЫЕ МАРКЕРЫ АССЕССОРА (исключаем) =====
        # Только ЯВНЫЕ признаки ассессора - прямые вопросы и обращения
        strong_assessor_markers = [
            'расскажите', 'опишите', 'объясните', 'покажите',
            'попробуйте', 'представьте', 'вообразите',
            'показывает участник', 'демонстрирует участник',
            'кандидат показывает', 'кандидат демонстрирует',
            'наблюдаю у участника', 'вижу у кандидата'
        ]
        
        has_strong_assessor = any(pattern in sent_lower for pattern in strong_assessor_markers)
        
        # ===== ШАГ 3: РОЛЕВЫЕ МАРКЕРЫ (исключаем) =====
        role_markers = [
            'директор:', 'менеджер:', 'клиент:', 'коллега:',
            'в роли директора', 'играю роль', 'я как директор', 'я как менеджер'
        ]
        
        has_role = any(pattern in sent_lower for pattern in role_markers)
        
        # ===== ЛОГИКА ПРИНЯТИЯ РЕШЕНИЯ =====
        
        # Если есть ролевые маркеры - всегда исключаем
        if has_role:
            continue
        
        # Если есть сильные маркеры ассессора - исключаем
        if has_strong_assessor:
            continue
        
        # Если есть сильные маркеры участника - ВКЛЮЧАЕМ (даже если есть "вы", "?" и т.д.)
        if has_strong_participant:
            participant_sentences.append(sent)
            continue
        
        # Если нет сильных маркеров - применяем мягкие критерии
        # Слабые маркеры ассессора (только если нет маркеров участника)
        weak_assessor_markers = [
            sent_lower.endswith('?'),  # вопрос в конце
            sent_lower.startswith('вы '),  # начинается с "вы"
            sent_lower.startswith('а вы '),
        ]
        
        if any(weak_assessor_markers):
            continue
        
        # Если дошли сюда - это скорее всего общее описание, берем
        participant_sentences.append(sent)
    
    result = '. '.join(participant_sentences)
    logger.info(f"СТРОГАЯ фильтрация речи: {len(text.split())} → {len(result.split())} слов")
    
    # Если после фильтрации слишком мало текста - предупреждаем
    if len(result.split()) < 200:
        logger.warning(f"⚠️ После фильтрации осталось мало текста участника: {len(result.split())} слов")
        logger.warning("Возможно, нужна ручная проверка транскрипта")
    
    return result


# ============================================================================
# ЛЕММАТИЗАЦИЯ
# ============================================================================

@lru_cache(maxsize=10000)
def lemmatize_word(word: str) -> str:
    """Лемматизация с кешированием"""
    if not HAS_PYMORPHY:
        return word.lower()
    try:
        return morph.parse(word)[0].normal_form
    except:
        return word.lower()


def lemmatize_text(text: str) -> List[str]:
    """Лемматизация текста"""
    words = re.findall(r'\b[а-яёА-ЯЁ]+\b', text.lower())
    return [lemmatize_word(w) for w in words if w not in STOP_WORDS]


# ============================================================================
# КРИТИЧЕСКИЕ ИНЦИДЕНТЫ (STAR)
# ============================================================================

def extract_critical_incidents(text: str, max_incidents: int = 5) -> List[Dict[str, Any]]:
    """
    Извлечение критических инцидентов (конкретных ситуаций по методу STAR)
    
    Args:
        text: очищенный текст
        max_incidents: максимальное количество инцидентов
        
    Returns:
        список инцидентов
    """
    # Маркеры конкретных ситуаций
    situation_markers = [
        r'один раз', r'был случай', r'например', r'в проекте',
        r'когда я работал', r'была ситуация', r'помню как',
        r'в прошлом году', r'на прошлой работе'
    ]
    
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.split()) > 5]
    
    incidents = []
    i = 0
    
    while i < len(sentences) and len(incidents) < max_incidents:
        sent = sentences[i]
        sent_lower = sent.lower()
        
        # Проверяем наличие маркера ситуации
        has_situation_marker = any(re.search(marker, sent_lower) for marker in situation_markers)
        
        if has_situation_marker:
            # Собираем контекст (текущее + следующие 2-3 предложения)
            context_sentences = sentences[i:min(i+4, len(sentences))]
            context = " ".join(context_sentences)
            
            # Пытаемся извлечь STAR компоненты
            incident = _extract_star_components(context, sent)
            
            if incident:
                incidents.append(incident)
                i += 3  # пропускаем обработанные предложения
            else:
                i += 1
        else:
            i += 1
    
    logger.info(f"Извлечено {len(incidents)} критических инцидентов")
    return incidents


def _extract_star_components(context: str, situation_sentence: str) -> Dict[str, Any]:
    """
    Извлечение компонентов STAR из контекста
    
    Returns:
        словарь с компонентами или None если не удалось извлечь
    """
    # Простая эвристика для извлечения действий (глаголы в прошедшем времени)
    action_verbs = []
    words = context.split()
    
    for word in words:
        lemma = lemmatize_word(word)
        if HAS_PYMORPHY:
            parsed = morph.parse(word)
            if parsed:
                tags = parsed[0].tag
                if 'VERB' in str(tags) and 'past' in str(tags):
                    action_verbs.append(lemma)
    
    if len(action_verbs) < 2:
        return None  # Слишком мало действий - вероятно не конкретная ситуация
    
    # Проверка на конкретность (упоминание людей, мест, проектов)
    has_specifics = bool(re.search(r'\b[А-ЯЁ][а-яё]+\b', context))  # Имена с заглавной
    
    confidence = "high" if (has_specifics and len(action_verbs) >= 3) else "medium"
    
    return {
        "situation": situation_sentence[:200],  # первые 200 символов
        "action": ", ".join(action_verbs[:5]),  # топ-5 действий
        "context_full": context[:500],  # полный контекст (макс 500 символов)
        "confidence": confidence,
        "action_count": len(action_verbs)
    }


# ============================================================================
# МОДУЛЬ 1: MBTI
# ============================================================================

def analyze_mbti(text: str, speech_metrics: Dict[str, Any] = None) -> Dict[str, Any]:
    """Анализ MBTI типа с учетом метрик речи"""
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    if total_words < 100:
        return {"confidence": "low", "note": "Недостаточно текста (минимум 100 слов)"}
    
    scores = {}
    
    # Анализ каждой дихотомии
    for dimension, markers in MBTI_MARKERS.items():
        pole_counts = {}
        for pole, marker_list in markers.items():
            count = sum(1 for lemma in lemmas if any(marker in lemma for marker in marker_list))
            pole_counts[pole] = count
        
        # Нормализация
        total = sum(pole_counts.values())
        if total > 0:
            scores[dimension] = {k: round(v / total, 2) for k, v in pole_counts.items()}
        else:
            scores[dimension] = {k: 0.5 for k in pole_counts.keys()}
    
    # КОРРЕКТИРОВКА ПО МЕТРИКАМ РЕЧИ (ПРИОРИТЕТ над лексическим анализом)
    if speech_metrics:
        wpm = speech_metrics.get('words_per_minute', 120)
        avg_utt = speech_metrics.get('avg_words_per_utterance', 10)
        
        # КРИТИЧЕСКИ ВАЖНО: короткие реплики (< 5 слов) = ИНТРОВЕРСИЯ
        # Интроверт экономит социальную энергию, говорит кратко
        if avg_utt < 5:
            # Сильная коррекция в сторону интроверсии
            scores["E_I"]["I"] = min(1.0, 0.75)  # Устанавливаем высокую интроверсию
            scores["E_I"]["E"] = max(0.0, 0.25)
            logger.info(f"✅ КРИТЕРИЙ: Короткие реплики ({avg_utt} слов) → ИНТРОВЕРСИЯ (I)")
        
        # Длинные развернутые реплики (> 10 слов) = ЭКСТРАВЕРСИЯ
        # Экстраверт получает энергию от общения, говорит развернуто
        elif avg_utt > 10:
            scores["E_I"]["E"] = min(1.0, scores["E_I"]["E"] + 0.20)
            scores["E_I"]["I"] = max(0.0, scores["E_I"]["I"] - 0.20)
            logger.info(f"✅ Развернутые реплики ({avg_utt} слов) → экстраверсия (E)")
    
    # Определение типа
    mbti_type = ""
    mbti_type += "E" if scores["E_I"]["E"] > scores["E_I"]["I"] else "I"
    mbti_type += "S" if scores["S_N"]["S"] > scores["S_N"]["N"] else "N"
    mbti_type += "T" if scores["T_F"]["T"] > scores["T_F"]["F"] else "F"
    mbti_type += "J" if scores["J_P"]["J"] > scores["J_P"]["P"] else "P"
    
    # Уверенность
    confidence = min(100, (total_words / 500) * 100)
    
    # Извлечение ключевых цитат
    key_quotes = _extract_quotes_for_dimension(text, mbti_type)
    
    return {
        "predicted_type": mbti_type,
        "confidence": round(confidence, 1),
        "E_I_score": scores["E_I"],
        "S_N_score": scores["S_N"],
        "T_F_score": scores["T_F"],
        "J_P_score": scores["J_P"],
        "key_quotes": key_quotes[:3]
    }


def _extract_quotes_for_dimension(text: str, mbti_type: str) -> List[str]:
    """Извлечение цитат, подтверждающих MBTI тип"""
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.split()) >= 5]
    
    quotes = []
    
    # Для каждой буквы типа ищем подтверждающие предложения
    for letter in mbti_type:
        for dimension, markers in MBTI_MARKERS.items():
            if letter in markers:
                marker_list = markers[letter]
                for sent in sentences:
                    if any(marker in sent.lower() for marker in marker_list[:3]):
                        if sent not in quotes:
                            quotes.append(sent)
                        if len(quotes) >= 5:
                            break
    
    return quotes


# ============================================================================
# МОДУЛЬ 2: ПАВЛОВ
# ============================================================================

def analyze_pavlov(text: str, speech_metrics: Dict[str, Any] = None) -> Dict[str, Any]:
    """Анализ типа нервной системы по Павлову с учетом метрик речи"""
    lemmas = lemmatize_text(text)
    text_lower = text.lower()
    total_words = len(lemmas)
    
    if total_words < 100:
        return {"confidence": "low", "note": "Недостаточно текста"}
    
    scores = {}
    
    for parameter, poles in PAVLOV_MARKERS.items():
        pole_scores = {}
        for pole_name, markers in poles.items():
            # Ищем и в леммах и в оригинальном тексте
            count_lemmas = sum(1 for lemma in lemmas if any(m in lemma for m in markers))
            count_text = sum(1 for m in markers if m in text_lower)
            count = max(count_lemmas, count_text)  # берем максимум
            pole_scores[pole_name] = count
        
        scores[parameter] = pole_scores
    
    # КОРРЕКТИРОВКА ПО МЕТРИКАМ РЕЧИ
    if speech_metrics:
        wpm = speech_metrics.get('words_per_minute', 120)
        
        # Подвижность КОРРЕЛИРУЕТ со скоростью речи (СТРОГИЕ пороги!)
        if wpm > 180:  # ДЕЙСТВИТЕЛЬНО быстрая речь
            scores["подвижность"]["mobile"] += 5
            logger.info(f"✅ Очень высокая скорость речи ({wpm} WPM) → +подвижность")
        elif wpm < 120:  # Медленная речь
            scores["подвижность"]["inert"] += 5
            logger.info(f"✅ Низкая скорость речи ({wpm} WPM) → +инертность")
    
    # Определение типа
    strength = "strong" if scores["сила"]["strong"] > scores["сила"]["weak"] else "weak"
    mobility = "mobile" if scores["подвижность"]["mobile"] > scores["подвижность"]["inert"] else "inert"
    
    balance_scores = scores["баланс"]
    balance = max(balance_scores, key=balance_scores.get)
    
    # ПРАВИЛЬНАЯ КЛАССИФИКАЦИЯ ПО ПАВЛОВУ:
    # 1. Меланхолик: СЛАБЫЙ (независимо от остального)
    # 2. Сангвиник: Сильный + Уравновешенный + Подвижный
    # 3. Флегматик: Сильный + Уравновешенный + Инертный
    # 4. Холерик: Сильный + Неуравновешенный (возбуждение преобладает)
    
    if strength == "weak":
        nervous_system_type = "Слабый (меланхолик)"
    elif strength == "strong":
        if balance == "balanced":
            # Уравновешенный
            if mobility == "mobile":
                nervous_system_type = "Сильный уравновешенный подвижный (сангвиник)"
            else:  # inert
                nervous_system_type = "Сильный уравновешенный инертный (флегматик)"
        elif balance == "excitation":
            # Неуравновешенный с преобладанием возбуждения
            nervous_system_type = "Сильный неуравновешенный (холерик)"
        elif balance == "inhibition":
            # Неуравновешенный с преобладанием торможения
            nervous_system_type = "Сильный с преобладанием торможения (сдержанный тип)"
        else:
            nervous_system_type = f"Сильный тип (требует уточнения баланса)"
    else:
        nervous_system_type = "Тип не определен"
    
    return {
        "nervous_system_type": nervous_system_type,
        "strength": strength,
        "mobility": mobility,
        "balance": balance,
        "scores": scores,
        "confidence": min(100, (total_words / 500) * 100)
    }


# ============================================================================
# МОДУЛЬ 3: OCEAN
# ============================================================================

def analyze_ocean(text: str, speech_metrics: Dict[str, Any] = None) -> Dict[str, Any]:
    """Анализ Большой пятерки (OCEAN) с учетом метрик речи"""
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    def calculate_trait_score(markers_dict):
        high_count = sum(1 for lemma in lemmas if any(m in lemma for m in markers_dict["high"]))
        low_count = sum(1 for lemma in lemmas if any(m in lemma for m in markers_dict["low"]))
        
        # Нормализация в шкалу 0-10 с учетом частоты
        total = high_count + low_count
        if total == 0:
            return 5.0  # нейтральное значение при отсутствии маркеров
        
        # Если маркеров мало (< 3) - ближе к середине
        if total < 3:
            ratio = high_count / total
            # Сжимаем к центру: 0→3, 0.5→5, 1→7
            score = 3 + (ratio * 4)
        else:
            # Нормальный расчет: диапазон 1-9 (избегаем только крайнего 10)
            ratio = high_count / total
            score = 1 + (ratio * 8)  # 0→1, 0.5→5, 1→9
        
        return round(score, 1)
    
    ocean_scores = {}
    for trait, markers in OCEAN_MARKERS.items():
        ocean_scores[trait] = calculate_trait_score(markers)
    
    # КОРРЕКТИРОВКА ПО МЕТРИКАМ РЕЧИ
    if speech_metrics:
        wpm = speech_metrics.get('words_per_minute', 120)

        # Extraversion (E/I) меряется по ВОКАЛЬНОЙ ЭНЕРГИИ (IQR громкости, дБ), а НЕ по
        # объёму/темпу речи: ровная, маловариативная подача → интроверсия; динамичная → экстраверсия.
        vocal_energy_iqr = speech_metrics.get('vocal_energy_iqr')
        if vocal_energy_iqr is not None:
            ocean_scores["Extraversion"] = round(
                max(1.0, min(9.0, 1.0 + (vocal_energy_iqr - 1.5) * 0.55)), 1)
            logger.info(f"✅ OCEAN: E/I по вокальной энергии (IQR={vocal_energy_iqr} дБ) → {ocean_scores['Extraversion']}")
        # Темп и объём речи на экстраверсию НЕ влияют (намеренно).

        # Neuroticism может проявляться в ЭКСТРЕМАЛЬНОЙ скорости
        if wpm and (wpm > 200 or wpm < 80):
            ocean_scores["Neuroticism"] = min(9.0, ocean_scores["Neuroticism"] + 0.5)
            logger.info(f"✅ OCEAN: Экстремальная скорость речи ({wpm} WPM) → +нейротизм")
    
    result = {
        "Openness": ocean_scores["Openness"],
        "Conscientiousness": ocean_scores["Conscientiousness"],
        "Extraversion": ocean_scores["Extraversion"],
        "Agreeableness": ocean_scores["Agreeableness"],
        "Neuroticism": ocean_scores["Neuroticism"],
        "confidence": min(100, (total_words / 500) * 100)
    }
    
    # Если мало текста - добавляем предупреждение
    if total_words < 100:
        result["note"] = "Недостаточно текста для надежной оценки"
    
    return result


# ============================================================================
# МОДУЛЬ 4: 7 РАДИКАЛОВ
# ============================================================================

def analyze_radicals(text: str, speech_metrics: Dict[str, Any] = None) -> Dict[str, Any]:
    """Анализ 7 радикалов личности с проверкой метрик речи"""
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    radical_scores = {}
    
    for radical, markers in RADICALS_MARKERS.items():
        count = sum(1 for lemma in lemmas if any(marker in lemma for marker in markers))
        radical_scores[radical] = count
    
    # Нормализация к 0-1
    total_markers = sum(radical_scores.values())
    if total_markers > 0:
        normalized_scores = {k: round(v / total_markers, 2) for k, v in radical_scores.items()}
    else:
        normalized_scores = {k: 0.14 for k in radical_scores.keys()}  # равномерное распределение
    
    # ПРОВЕРКА ПРОТИВОРЕЧИЙ С МЕТРИКАМИ РЕЧИ
    if speech_metrics:
        wpm = speech_metrics.get('words_per_minute', 120)
        avg_utterance = speech_metrics.get('avg_words_per_utterance', 10)
        
        logger.info(f"🔍 Проверка радикалов с метриками: WPM={wpm}, Avg={avg_utterance}")
        logger.info(f"   Гипертимный ДО коррекции: {normalized_scores.get('гипертимный', 0)}")
        
        # Гипертимный радикал требует ВЫСОКОЙ скорости речи (>170 WPM)
        if normalized_scores.get('гипертимный', 0) > 0.10:
            if wpm < 170 or avg_utterance > 10:
                # Медленная речь или длинные реплики = УБИРАЕМ гипертимность ПОЛНОСТЬЮ
                logger.info(f"   ⚠️ КОРРЕКЦИЯ: НЕбыстрая речь ({wpm} WPM, avg={avg_utterance}) → УБИРАЕМ гипертимность")
                removed = normalized_scores['гипертимный']  # убираем ВСЁ!
                normalized_scores['гипертимный'] = 0.0  # ставим 0
                
                # Перераспределяем между флегматичными/инертными радикалами
                other_radicals = ['эпилептоидный', 'тревожный', 'шизоидный']
                for rad in other_radicals:
                    normalized_scores[rad] = normalized_scores.get(rad, 0) + (removed / len(other_radicals))
        
        # Тревожный радикал увеличивается при медленной речи
        if wpm < 100:
            logger.info(f"   ⚠️ Медленная речь ({wpm} WPM) → повышение тревожности")
            # Перераспределяем в пользу тревожного
            if normalized_scores.get('гипертимный', 0) > 0.1:
                transfer = normalized_scores['гипертимный'] * 0.2
                normalized_scores['гипертимный'] -= transfer
                normalized_scores['тревожный'] = normalized_scores.get('тревожный', 0) + transfer
        
        # Пересчет после корректировки (нормализация к 1.0)
        total_adjusted = sum(normalized_scores.values())
        if total_adjusted > 0:
            normalized_scores = {k: round(v / total_adjusted, 2) for k, v in normalized_scores.items()}
    
    # Определение доминантных радикалов (топ-2)
    sorted_radicals = sorted(normalized_scores.items(), key=lambda x: x[1], reverse=True)
    dominant = [r[0] for r in sorted_radicals[:2]]
    
    result = {
        "scores": normalized_scores,
        "dominant_radicals": dominant,
        "confidence": min(100, (total_words / 500) * 100)
    }
    
    if total_words < 100:
        result["note"] = "Недостаточно текста для надежной оценки"
    
    return result


# ============================================================================
# МОДУЛЬ 5: ТЕМНАЯ ТЕТРАДА
# ============================================================================

def analyze_dark_tetrad(text: str) -> Dict[str, Any]:
    """
    Анализ темной тетрады
    ОСТОРОЖНО: деликатная зона, используем мягкие интерпретации
    """
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    # Подсчет маркеров для каждой черты
    results = {}
    
    for trait, data in DARK_TETRAD_MARKERS.items():
        indicators = data["indicators"]
        count = sum(1 for lemma in lemmas if any(ind in lemma for ind in indicators))
        
        # Нормализация
        score = min(1.0, count / (total_words / 100)) if total_words > 0 else 0.0  # на 100 слов
        
        # Категоризация
        if score < 0.3:
            level = "low"
            interpretation = "Адаптивные черты в пределах нормы"
        elif score < 0.6:
            level = "medium"
            interpretation = "Умеренное проявление, требует контекстного анализа"
        else:
            level = "high"
            interpretation = "Повышенные показатели, рекомендуется дополнительная оценка"
        
        results[trait] = {
            "score": round(score, 2),
            "level": level,
            "interpretation": interpretation
        }
    
    # Дополнительно: анализ соотношения я/мы для нарциссизма
    i_count = text.lower().count(' я ')
    we_count = text.lower().count(' мы ')
    i_we_ratio = round(i_count / we_count, 2) if we_count > 0 else float('inf')
    
    results["narcissism"]["i_we_ratio"] = i_we_ratio
    if i_we_ratio > 2.0:
        results["narcissism"]["note"] = "Высокое соотношение 'я'/'мы' - возможна эгоцентричность"
    
    result = {
        "assessments": results,
        "confidence": min(100, (total_words / 500) * 100),
        "disclaimer": "Оценки являются индикативными и требуют профессиональной интерпретации"
    }
    
    if total_words < 100:
        result["note"] = "Недостаточно текста для надежной оценки"
    
    return result


# ============================================================================
# МОДУЛЬ 6: ПОТЕНЦИАЛ (Обучаемость)
# ============================================================================

def analyze_potential(text: str) -> Dict[str, Any]:
    """Анализ потенциала и обучаемости"""
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    # Подсчет маркеров для каждого аспекта
    results = {}
    
    for aspect, markers in POTENTIAL_MARKERS.items():
        high_count = sum(1 for lemma in lemmas if any(m in lemma for m in markers["high"]))
        low_count = sum(1 for lemma in lemmas if any(m in lemma for m in markers["low"]))
        
        # Нормализация с избежанием крайностей
        total = high_count + low_count
        if total > 0:
            ratio = high_count / total
            # Избегаем крайностей: 0→0.1, 0.5→0.5, 1→0.9
            if total < 3:
                # Мало данных - сжимаем к центру
                score = 0.3 + (ratio * 0.4)  # диапазон 0.3-0.7
            else:
                # Достаточно данных - диапазон 0.1-0.9 (10%-90%)
                score = 0.1 + (ratio * 0.8)  # 0→0.1, 0.5→0.5, 1→0.9
        else:
            score = 0.5  # нейтрально если нет данных
        
        results[aspect] = {
            "score": round(score, 2),
            "high_indicators": high_count,
            "low_indicators": low_count
        }
    
    # Общий уровень потенциала
    avg_score = sum(r["score"] for r in results.values()) / len(results)
    
    if avg_score > 0.7:
        level = "высокий"
    elif avg_score > 0.5:
        level = "средний"
    else:
        level = "требует поддержки"
    
    result = {
        "level": level,
        "overall_score": round(avg_score, 2),
        "components": results,
        "confidence": min(100, (total_words / 500) * 100)
    }
    
    if total_words < 100:
        result["note"] = "Недостаточно текста для надежной оценки"
    
    return result


# ============================================================================
# МОДУЛЬ 7: СТРЕСС (Стрессоустойчивость)
# ============================================================================

def analyze_stress(text: str, disfluency_profile: Dict = None, speech_metrics: Dict[str, Any] = None) -> Dict[str, Any]:
    """
    Анализ стрессоустойчивости с учетом метрик речи
    Использует кросс-валидацию с дефлюенси + темп речи
    """
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    # Подсчет стрессоров
    stressor_count = sum(1 for lemma in lemmas if any(s in lemma for s in STRESS_MARKERS["stressors"]))
    
    # Adaptive coping
    adaptive_count = sum(1 for lemma in lemmas if any(c in lemma for c in STRESS_MARKERS["adaptive_coping"]))
    
    # Maladaptive coping
    maladaptive_count = sum(1 for lemma in lemmas if any(c in lemma for c in STRESS_MARKERS["maladaptive_coping"]))
    
    # Recovery
    fast_recovery = sum(1 for lemma in lemmas if any(r in lemma for r in STRESS_MARKERS["recovery"]["fast"]))
    slow_recovery = sum(1 for lemma in lemmas if any(r in lemma for r in STRESS_MARKERS["recovery"]["slow"]))
    
    # Balance score
    positive_coping = adaptive_count + fast_recovery
    negative_coping = maladaptive_count + slow_recovery + (stressor_count * 0.5)
    
    balance = positive_coping - negative_coping
    
    # Определение уровня
    if balance > 5:
        level = "высокий"
    elif balance > 0:
        level = "средний"
    else:
        level = "требует внимания"
    
    result = {
        "level": level,
        "balance_score": round(balance, 2),
        "stressors_count": stressor_count,
        "adaptive_coping_count": adaptive_count,
        "maladaptive_coping_count": maladaptive_count,
        "recovery_speed": "fast" if fast_recovery > slow_recovery else "slow",
        "confidence": min(100, (total_words / 500) * 100)
    }
    
    # Кросс-валидация с дефлюенси если доступно
    if disfluency_profile:
        baseline_filler = disfluency_profile.get("baseline_filler_rate", 0)
        if baseline_filler > 0.15:
            result["disfluency_note"] = "Высокая дефлюентность может указывать на фоновый стресс"
    
    # КОРРЕКТИРОВКА ПО МЕТРИКАМ РЕЧИ
    if speech_metrics:
        wpm = speech_metrics.get('words_per_minute', 120)
        
        # Очень быстрая речь (>160) или очень медленная (<80) = признак стресса
        if wpm > 160:
            balance -= 2
            result["speech_note"] = "Очень быстрая речь может указывать на нервозность"
            logger.info(f"⚠️ Очень быстрая речь ({wpm} WPM) → повышенный стресс")
        elif wpm < 80:
            balance -= 2
            result["speech_note"] = "Очень медленная речь может указывать на подавленность"
            logger.info(f"⚠️ Очень медленная речь ({wpm} WPM) → повышенный стресс")
        
        # Пересчитываем уровень с учетом коррекции
        if balance > 5:
            level = "высокий"
        elif balance > 0:
            level = "средний"
        else:
            level = "требует внимания"
        
        result["level"] = level
        result["balance_score"] = round(balance, 2)
    
    if total_words < 100:
        result["note"] = "Недостаточно текста для надежной оценки"
    
    return result


# ============================================================================
# МОДУЛЬ 8: ТЕНЕВЫЕ СТОРОНЫ
# ============================================================================

def analyze_shadow_sides(text: str, strengths: List[str] = None) -> Dict[str, Any]:
    """
    Анализ теневых сторон (риски при усилении сильных сторон)
    """
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    if total_words < 100:
        return {"confidence": "low", "note": "Недостаточно текста"}
    
    identified_shadows = []
    
    # Анализируем каждую пару сила-тень
    for shadow_name, data in SHADOW_LOGIC.items():
        strength_markers = data["strength"]
        shadow_markers = data["shadow"]
        
        # Подсчет маркеров силы
        strength_count = sum(1 for lemma in lemmas if any(m in lemma for m in strength_markers))
        
        # Подсчет маркеров тени
        shadow_count = sum(1 for lemma in lemmas if any(m in lemma for m in shadow_markers))
        
        # Если есть и сила и тень - это значимо
        if strength_count > 0 and shadow_count > 0:
            risk_level = "high" if shadow_count >= strength_count else "medium"
            
            identified_shadows.append({
                "strength": shadow_name,
                "strength_indicators": strength_count,
                "shadow_indicators": shadow_count,
                "risk_level": risk_level
            })
    
    # Сортируем по риску
    identified_shadows.sort(key=lambda x: x["shadow_indicators"], reverse=True)
    
    return {
        "identified_shadows": identified_shadows[:3],  # топ-3
        "total_analyzed": len(SHADOW_LOGIC),
        "confidence": min(100, (total_words / 500) * 100)
    }


# ============================================================================
# МОДУЛЬ 9: PRISM (5 слоев личности)
# ============================================================================

def analyze_prism(text: str, disfluency_profile: Dict = None) -> Dict[str, Any]:
    """
    PRISM анализ - 5 слоев личности
    Использует кросс-валидацию с дефлюенси для слоя Resistance
    """
    lemmas = lemmatize_text(text)
    total_words = len(lemmas)
    
    if total_words < 100:
        return {"confidence": "low", "note": "Недостаточно текста"}
    
    layers = {}
    
    # Анализ каждого слоя
    for layer_name, data in PRISM_MARKERS.items():
        indicators = data["indicators"]
        
        # Ищем предложения с маркерами этого слоя
        sentences = re.split(r'[.!?]+', text)
        layer_sentences = []
        
        for sent in sentences:
            if any(ind in sent.lower() for ind in indicators):
                layer_sentences.append(sent.strip())
        
        layers[layer_name] = {
            "found_indicators": len(layer_sentences),
            "examples": layer_sentences[:3]  # топ-3 примера
        }
    
    # Специальная обработка слоя Resistance с дефлюенси
    if disfluency_profile and "Resistance" in layers:
        baseline = disfluency_profile.get("baseline_filler_rate", 0)
        if baseline > 0.15:
            layers["Resistance"]["disfluency_correlation"] = "High baseline disfluency indicates general discomfort zones"
    
    return {
        "layers": layers,
        "confidence": min(100, (total_words / 500) * 100)
    }


# ============================================================================
# МОДУЛЬ 10: КОМПЕТЕНЦИИ (с чтением МК.xlsx)
# ============================================================================

def analyze_competencies(text: str, competency_file: str = None) -> Dict[str, Any]:
    """
    Анализ по компетенциям из модели компании
    """
    lemmas = lemmatize_text(text)
    text_lower = text.lower()
    total_words = len(lemmas)
    
    if total_words < 100:
        return {"confidence": "low", "note": "Недостаточно текста"}
    
    # Базовые компетенции с расширенными маркерами
    default_competencies = {
        "Лидерство": ["лидер", "возглавл", "руковод", "организ", "инициатив", "веду", "направля"],
        "Коммуникация": ["общ", "говор", "объясн", "презент", "обсужд", "договор", "коммуник"],
        "Аналитическое мышление": ["анализ", "анализир", "данн", "логик", "систем", "структур", "разбир"],
        "Командная работа": ["команд", "вместе", "сотрудн", "коллектив", "группов", "совместн"],
        "Стратегическое мышление": ["стратег", "долгосрочн", "план", "перспектив", "цель", "видение"],
        "Управление изменениями": ["изменен", "адаптац", "гибк", "нов", "трансформ", "перестра"],
        "Принятие решений": ["решен", "решил", "выбор", "определ", "принял", "принима"],
        "Развитие подчиненных": ["обуч", "развив", "наставн", "ментор", "помог", "учил"],
        "Целостность личности": ["надежн", "честн", "контрол", "уверен", "эмоц"],
        "Ориентация на результат": ["результат", "достиж", "цель", "эффективн", "качеств"]
    }
    
    # Попытка загрузить МК.xlsx с индикаторами
    competencies = {}
    if competency_file and os.path.exists(competency_file):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(competency_file)
            ws = wb.active
            
            # Первая колонка - название, вторая - индикаторы
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row[0] and row[1]:
                    comp_name = str(row[0]).strip()
                    indicators_text = str(row[1]).lower()
                    
                    # Извлекаем ключевые слова из индикаторов (простая эвристика)
                    key_words = []
                    
                    # Берем важные слова из описания
                    important_words = re.findall(r'\b[а-яё]{5,}\b', indicators_text)
                    # Лемматизируем
                    key_words = [lemmatize_word(w) for w in important_words[:30]]  # топ-30 слов
                    
                    competencies[comp_name] = key_words
            
            if competencies:
                logger.info(f"Загружено {len(competencies)} компетенций с индикаторами из {competency_file}")
            else:
                competencies = default_competencies
        except Exception as e:
            logger.warning(f"Не удалось загрузить {competency_file}: {e}. Использую базовые компетенции")
            competencies = default_competencies
    else:
        competencies = default_competencies
    
    # Анализ каждой компетенции
    results = {}
    
    for comp_name, markers in competencies.items():
        # Подсчет упоминаний - ищем и в леммах и в тексте
        count_lemmas = 0
        count_text = 0
        
        for lemma in lemmas:
            if any(marker in lemma for marker in markers):
                count_lemmas += 1
        
        for marker in markers:
            if marker in text_lower:
                count_text += 1
        
        # Берем максимум из двух методов
        count = max(count_lemmas, count_text)
        
        # Контекстный анализ (упрощенный)
        sentences = re.split(r'[.!?]+', text)
        active_contexts = 0
        passive_contexts = 0
        
        for sent in sentences:
            sent_lower = sent.lower()
            if any(marker in sent_lower for marker in markers):
                # Простая эвристика: есть ли "я" в предложении
                if ' я ' in sent_lower:
                    active_contexts += 1
                else:
                    passive_contexts += 1
        
        # Depth score (насколько детально описана)
        depth = min(1.0, count / 3) if count > 0 else 0.0
        
        results[comp_name] = {
            "frequency": count,
            "contexts": {
                "active": active_contexts,
                "passive": passive_contexts
            },
            "depth_score": round(depth, 2),
            "present": count > 0
        }
    
    return {
        "competencies": results,
        "total_analyzed": len(competencies),
        "confidence": min(100, (total_words / 500) * 100)
    }


# ============================================================================
# МОДУЛЬ 10Б: КОМПЕТЕНЦИИ СЕМАНТИЧЕСКИЕ (с AI)
# ============================================================================

def analyze_competencies_semantic(text: str, competency_file: str = None) -> Dict[str, Any]:
    """
    Семантический анализ компетенций через Sentence-BERT
    НАМНОГО ТОЧНЕЕ чем поиск подстрок!
    """
    if not HAS_SEMANTIC:
        # Fallback на обычный анализ
        return analyze_competencies(text, competency_file)
    
    # Загружаем компетенции из Excel
    competencies_with_indicators = {}
    
    if competency_file and os.path.exists(competency_file):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(competency_file)
            ws = wb.active
            
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row[0] and row[1]:
                    comp_name = str(row[0]).strip()
                    indicators = str(row[1])
                    competencies_with_indicators[comp_name] = indicators
            
            logger.info(f"Загружено {len(competencies_with_indicators)} компетенций для семантического анализа")
        except Exception as e:
            logger.error(f"Ошибка загрузки МК.xlsx: {e}")
            return {"error": str(e)}
    else:
        # Базовые компетенции
        competencies_with_indicators = {
            "Лидерство": "организация команды принятие решений ответственность инициатива",
            "Коммуникация": "общение объяснение презентация обсуждение договоренности",
            "Аналитическое мышление": "анализ данные логика система структура",
            "Командная работа": "команда сотрудничество вместе коллектив группа"
        }
    
    # Разбиваем текст на предложения
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    if not sentences:
        return {"competencies": {}, "note": "Нет предложений для анализа"}
    
    # Получаем эмбеддинги предложений
    sentence_embeddings = semantic_model.encode(sentences)
    
    results = {}
    
    for comp_name, indicators_text in competencies_with_indicators.items():
        # Эмбеддинг индикаторов компетенции
        indicator_embedding = semantic_model.encode([indicators_text])[0]
        
        # Находим релевантные предложения
        similarities = []
        relevant_sentences = []
        
        for i, sent_emb in enumerate(sentence_embeddings):
            similarity = np.dot(indicator_embedding, sent_emb) / (
                np.linalg.norm(indicator_embedding) * np.linalg.norm(sent_emb)
            )
            
            if similarity > 0.2:  # порог релевантности (понижен с 0.25)
                similarities.append(similarity)
                relevant_sentences.append({
                    "text": sentences[i],
                    "score": float(similarity)
                })
        
        # Сортируем по релевантности
        relevant_sentences.sort(key=lambda x: x["score"], reverse=True)
        
        # Определяем active/passive контекст
        active = sum(1 for s in relevant_sentences if ' я ' in s["text"].lower())
        passive = len(relevant_sentences) - active
        
        results[comp_name] = {
            "frequency": len(similarities),
            "relevance_score": round(float(np.mean(similarities)), 3) if similarities else 0.0,
            "contexts": {
                "active": active,
                "passive": passive
            },
            "depth_score": min(1.0, len(similarities) / 3),
            "top_examples": [s["text"][:150] for s in relevant_sentences[:3]],
            "present": len(similarities) > 0
        }
    
    return {
        "competencies": results,
        "total_analyzed": len(competencies_with_indicators),
        "method": "semantic_ai",
        "confidence": 90  # семантический анализ более надежен
    }


# ============================================================================
# МОДУЛЬ 11: КРАСНАЯ НИТЬ (Кросс-валидация)
# ============================================================================

def cross_validate_exercises(exercises_data: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Кросс-валидация паттернов через все упражнения
    Ищет стабильные характеристики, присутствующие везде
    """
    if len(exercises_data) < 2:
        return {
            "note": "Требуется минимум 2 упражнения для кросс-валидации",
            "stable_patterns": [],
            "situational_patterns": []
        }
    
    # Здесь должна быть сложная логика сравнения паттернов
    # Упрощенная версия для MVP
    
    stable_patterns = []
    situational_patterns = []
    
    # Пример: если MBTI тип одинаковый во всех упражнениях = стабильный паттерн
    mbti_types = [ex.get("mbti_markers", {}).get("predicted_type") for ex in exercises_data]
    if len(set(mbti_types)) == 1 and mbti_types[0]:
        stable_patterns.append({
            "trait": f"MBTI тип: {mbti_types[0]}",
            "stability": "stable",
            "evidence_count": len(exercises_data)
        })
    
    return {
        "stable_patterns": stable_patterns,
        "situational_patterns": situational_patterns,
        "exercises_analyzed": len(exercises_data)
    }


# ============================================================================
# МОДУЛЬ 12: УПРАВЛЕНЧЕСКИЕ РЕКОМЕНДАЦИИ (Синтез)
# ============================================================================

def prepare_management_data(analysis_result: Dict[str, Any]) -> Dict[str, Any]:
    """
    Подготовка данных для управленческих рекомендаций
    Синтез из всех модулей
    """
    
    # Извлекаем ключевые данные из всех модулей
    mbti = analysis_result.get("mbti_markers", {})
    ocean = analysis_result.get("ocean_markers", {})
    stress = analysis_result.get("stress_profile", {})
    potential = analysis_result.get("potential_markers", {})
    
    # Мотиваторы (на основе типологий)
    motivation_drivers = []
    mbti_type = mbti.get("predicted_type", "")
    if "N" in mbti_type:
        motivation_drivers.append("Интеллектуальные вызовы и новые идеи")
    if "J" in mbti_type:
        motivation_drivers.append("Структура и ясные цели")
    if "E" in mbti_type:
        motivation_drivers.append("Командная работа и коммуникация")
    
    # Стрессоры
    stress_triggers = []
    if stress.get("level") == "требует внимания":
        stress_triggers.append("Высокая нагрузка без поддержки")
    
    # Предпочтения в коммуникации
    communication_preferences = []
    if "I" in mbti_type:
        communication_preferences.append("Письменная коммуникация или время на подготовку")
    
    # Зоны развития
    development_focus = []
    if potential.get("level") in ["средний", "требует поддержки"]:
        development_focus.append("Обучаемость и адаптивность")
    
    # Сильные стороны
    strengths = []
    if ocean.get("Conscientiousness", 5) > 7:
        strengths.append("Организованность и надежность")
    if ocean.get("Openness", 5) > 7:
        strengths.append("Открытость новому опыту")
    
    return {
        "motivation_drivers": motivation_drivers or ["Требуется дополнительный анализ"],
        "communication_preferences": communication_preferences or ["Стандартный подход"],
        "stress_triggers": stress_triggers or ["Стандартные стрессоры"],
        "development_focus": development_focus or ["Определяется индивидуально"],
        "strengths_to_leverage": strengths or ["Требуется углубленный анализ"]
    }


# ============================================================================
# ПСИХОЛИНГВИСТИЧЕСКИЙ АНАЛИЗ (Расширенный)
# ============================================================================

def analyze_psycholinguistics(text: str) -> Dict[str, Any]:
    """Расширенный психолингвистический анализ"""
    words = text.lower().split()
    total_words = len(words)
    
    if total_words < 50:
        return {"confidence": "low", "note": "Недостаточно текста"}
    
    # 1. Модальность
    confident_verbs = len(re.findall(r'\b(сделал|организовал|решил|выбрал|достиг)\b', text.lower()))
    uncertain_verbs = len(re.findall(r'\b(попробовал|попытался|постарался|вроде)\b', text.lower()))
    obligatory = len(re.findall(r'\b(нужно было|следовало|должен был)\b', text.lower()))
    
    total_modal = confident_verbs + uncertain_verbs + obligatory
    if total_modal > 0:
        modality_profile = {
            "confident": round(confident_verbs / total_modal, 2),
            "uncertain": round(uncertain_verbs / total_modal, 2),
            "obligatory": round(obligatory / total_modal, 2)
        }
    else:
        modality_profile = {"confident": 0.5, "uncertain": 0.3, "obligatory": 0.2}
    
    # 2. Локус контроля
    active_phrases = len(re.findall(r'\bя (решил|сделал|выбрал|организовал)\b', text.lower()))
    passive_phrases = len(re.findall(r'\b(меня попросили|пришлось|обстоятельства)\b', text.lower()))
    
    if active_phrases + passive_phrases > 0:
        locus = "internal" if active_phrases > passive_phrases else "external"
    else:
        locus = "mixed"
    
    # 3. Я vs МЫ
    i_count = text.lower().count(' я ')
    we_count = text.lower().count(' мы ')
    i_we_ratio = round(i_count / we_count, 2) if we_count > 0 else 0
    
    if i_we_ratio > 2:
        interpretation = "Сильный индивидуальный фокус"
    elif i_we_ratio > 1:
        interpretation = "Баланс с легким индивидуальным акцентом"
    elif i_we_ratio > 0.5:
        interpretation = "Баланс между я и командой"
    else:
        interpretation = "Сильный командный фокус"
    
    return {
        "modality_profile": modality_profile,
        "locus_of_control": locus,
        "pronouns_ratio": {
            "i_count": i_count,
            "we_count": we_count,
            "ratio": i_we_ratio,
            "interpretation": interpretation
        },
        "confidence": min(100, (total_words / 300) * 100)
    }


# ============================================================================
# ГЛАВНАЯ ФУНКЦИЯ ОБРАБОТКИ
# ============================================================================

def process_exercises(exercises: List[str], participant_name: str = "Участник", speech_metrics: List[Dict[str, Any]] = None) -> Dict[str, Any]:
    """
    Полная обработка упражнений участника
    
    Args:
        exercises: список текстов упражнений
        participant_name: имя участника
        speech_metrics: метрики речи из аудио (если есть)
        
    Returns:
        компактная JSON структура для Claude API
    """
    logger.info(f"Начало обработки {len(exercises)} упражнений для {participant_name}")
    
    # Агрегируем метрики речи если есть
    aggregated_metrics = None
    if speech_metrics:
        total_wpm = sum(m.get('words_per_minute', 0) for m in speech_metrics)
        total_avg_utt = sum(m.get('avg_words_per_utterance', 0) for m in speech_metrics)
        count = len(speech_metrics)
        
        aggregated_metrics = {
            'words_per_minute': total_wpm / count if count > 0 else 120,
            'avg_words_per_utterance': total_avg_utt / count if count > 0 else 10,
            'has_audio_data': True
        }
        logger.info(f"✅ Используем агрегированные метрики речи: WPM={aggregated_metrics['words_per_minute']:.1f}, Avg={aggregated_metrics['avg_words_per_utterance']:.1f}")
    else:
        logger.info("⚠️ Метрики речи отсутствуют - анализ только по тексту")
    
    # Объединяем все упражнения
    all_text = "\n\n".join(exercises)
    
    # ЭТАП 0.5: Извлечение обсуждения ассессоров (если есть)
    logger.info("Этап 0.5: Поиск обсуждения ассессоров...")
    all_text, assessor_discussion = extract_assessor_discussion(all_text)
    
    # ЭТАП 1: Анализ дефлюенси НА СЫРОМ ТЕКСТЕ
    logger.info("Этап 1: Анализ дефлюенси...")
    disfluency_profile = analyze_disfluencies(all_text)
    
    # ЭТАП 2: Удаление легенды упражнения
    logger.info("Этап 2: Удаление легенды упражнения...")
    text_without_legend = remove_exercise_legend(all_text)
    
    # ЭТАП 2.5: Очистка текста
    logger.info("Этап 2.5: Очистка текста...")
    clean = clean_text(text_without_legend)
    
    # ЭТАП 2.7: Фильтрация речи участника (убираем речь ассессора)
    logger.info("Этап 2.7: Фильтрация речи участника...")
    participant_text = extract_participant_speech(clean)
    
    # Используем отфильтрованный текст для анализа
    clean = participant_text if len(participant_text) > 100 else clean
    
    # ЭТАП 3: Критические инциденты + выжимка участника
    logger.info("Этап 3: Извлечение критических инцидентов...")
    critical_incidents = extract_critical_incidents(clean, max_incidents=15)  # ✅ Увеличено с 5 до 15!
    
    # ЭТАП 4: Анализ по ВСЕМ 12 модулям
    logger.info("Этап 4: Анализ по психологическим модулям...")
    
    # Модули 1-5: базовые личностные типологии
    mbti_result = analyze_mbti(clean, speech_metrics=aggregated_metrics)  # ✅ МЕТРИКИ!
    pavlov_result = analyze_pavlov(clean, speech_metrics=aggregated_metrics)  # ✅ МЕТРИКИ!
    ocean_result = analyze_ocean(clean, speech_metrics=aggregated_metrics)  # ✅ МЕТРИКИ!
    radicals_result = analyze_radicals(clean, speech_metrics=aggregated_metrics)  # ✅ МЕТРИКИ!
    dark_tetrad_result = analyze_dark_tetrad(clean)
    
    # Модули 6-7: потенциал и стресс
    potential_result = analyze_potential(clean)
    stress_result = analyze_stress(clean, disfluency_profile, speech_metrics=aggregated_metrics)  # ✅ МЕТРИКИ!
    
    # Модули 8-9: теневые стороны и PRISM
    shadow_result = analyze_shadow_sides(clean)
    prism_result = analyze_prism(clean, disfluency_profile)
    
    # Модуль 10: компетенции (с семантическим анализом если доступен)
    competency_file = "config/МК.xlsx" if os.path.exists("config/МК.xlsx") else None
    
    if HAS_SEMANTIC:
        logger.info("Используем СЕМАНТИЧЕСКИЙ анализ компетенций (AI)")
        competencies_result = analyze_competencies_semantic(clean, competency_file)
    else:
        logger.info("Используем базовый анализ компетенций (regex)")
        competencies_result = analyze_competencies(clean, competency_file)
    
    # Дополнительно: психолингвистика
    psycholing_result = analyze_psycholinguistics(clean)
    
    # ЭТАП 5: Сборка промежуточной структуры для кросс-валидации
    logger.info("Этап 5: Кросс-валидация и синтез...")
    
    intermediate_result = {
        "mbti_markers": mbti_result,
        "pavlov_markers": pavlov_result,
        "ocean_markers": ocean_result,
        "radicals_markers": radicals_result,
        "dark_tetrad_markers": dark_tetrad_result,
        "potential_markers": potential_result,
        "stress_profile": stress_result,
        "shadow_sides": shadow_result,
        "prism_layers": prism_result,
        "competencies_analysis": competencies_result
    }
    
    # КРОСС-ВАЛИДАЦИЯ: Проверка противоречий между модулями
    contradictions = []
    
    # Противоречие 1: Флегматик (инертная НС) vs Открытость новому
    if pavlov_result.get('temperament') == 'Флегматик' or 'инертн' in pavlov_result.get('nervous_system_type', '').lower():
        openness = ocean_result.get('Openness', 5)
        if openness > 7:
            contradictions.append("Инертная НС (флегматик) + высокая Открытость новому (OCEAN) - ПРОТИВОРЕЧИЕ!")
            ocean_result['Openness'] = min(6.0, openness)  # Снижаем до реалистичного
            logger.warning(f"⚠️ ПРОТИВОРЕЧИЕ: Флегматик не может иметь Openness={openness}, снижено до {ocean_result['Openness']}")
    
    # Противоречие 2: Гипертимный радикал vs Интроверсия
    if radicals_result.get('scores', {}).get('гипертимный', 0) > 0.15:
        mbti_type = mbti_result.get('predicted_type', 'XXXX')
        if mbti_type.startswith('I'):
            contradictions.append("Гипертимный радикал + Интроверсия (MBTI) - ПРОТИВОРЕЧИЕ!")
            # Снижаем гипертимность в scores
            if 'scores' in radicals_result:
                radicals_result['scores']['гипертимный'] = 0.05
            logger.warning(f"⚠️ ПРОТИВОРЕЧИЕ: Интроверт не может быть гипертимным!")
    
    # Противоречие 3: Ригидность (эпилептоидный + инертный) vs Управление изменениями
    # КРИТИЧНО: Эпилептоид (ригидность) И/ИЛИ Флегматик (инертность) → плохое управление изменениями
    is_rigid = radicals_result.get('scores', {}).get('эпилептоидный', 0) > 0.20
    is_inert = pavlov_result.get('temperament') == 'Флегматик' or 'инертн' in pavlov_result.get('nervous_system_type', '').lower()
    
    if is_rigid or is_inert:
        # Проверяем компетенцию "управление изменениями"
        competencies_dict = competencies_result.get('competencies', {})
        if isinstance(competencies_dict, dict):
            for comp_name, comp_data in competencies_dict.items():
                if 'изменени' in comp_name.lower():
                    old_score = comp_data.get('relevance_score', 0)
                    if old_score > 0.35:  # Если оценка выше "Открыт для развития"
                        contradictions.append(f"{'Эпилептоид' if is_rigid else 'Флегматик'} + высокое Управление изменениями - ПРОТИВОРЕЧИЕ!")
                        # ЖЕСТКО снижаем до максимум 0.3 (соответствует оценке 2)
                        comp_data['relevance_score'] = min(0.30, old_score)
                        logger.warning(f"⚠️ ПРОТИВОРЕЧИЕ: Ригидность/инертность → Управление изменениями снижено с {old_score:.2f} до {comp_data['relevance_score']:.2f}")
    
    # Противоречие 4: Высокий нейротизм vs Целостность личности
    # КРИТИЧНО: Нейротизм > 7.0 → эмоциональный контроль нестабилен → Целостность максимум 3
    neuroticism = ocean_result.get('Neuroticism', 5)
    stress_balance = stress_result.get('balance', 0) if isinstance(stress_result, dict) else 0
    
    if neuroticism > 7.0 or stress_balance < -20:
        competencies_dict = competencies_result.get('competencies', {})
        if isinstance(competencies_dict, dict):
            for comp_name, comp_data in competencies_dict.items():
                if 'целостн' in comp_name.lower() or 'integrity' in comp_name.lower():
                    old_score = comp_data.get('relevance_score', 0)
                    if old_score > 0.50:  # Если оценка 4+
                        reason = f"Нейротизм={neuroticism:.1f}" if neuroticism > 7 else f"Стресс-баланс={stress_balance}"
                        contradictions.append(f"{reason} → Целостность не может быть 4+ (нестабильный эмоциональный контроль)")
                        # Снижаем до максимум 0.45 (соответствует оценке 3)
                        comp_data['relevance_score'] = min(0.45, old_score)
                        logger.warning(f"⚠️ ПРОТИВОРЕЧИЕ: {reason} → Целостность снижена с {old_score:.2f} до {comp_data['relevance_score']:.2f}")
    
    if contradictions:
        intermediate_result['contradictions_found'] = contradictions
        logger.info(f"🔍 Найдено и исправлено противоречий: {len(contradictions)}")
    
    # Модуль 11: Красная нить (если несколько упражнений)
    if len(exercises) > 1:
        # Для каждого упражнения делаем отдельный анализ
        exercises_data = []
        for ex in exercises:
            ex_mbti = analyze_mbti(ex)
            exercises_data.append({"mbti_markers": ex_mbti})
        red_thread_result = cross_validate_exercises(exercises_data)
    else:
        red_thread_result = {
            "note": "Одно упражнение - кросс-валидация недоступна",
            "stable_patterns": [],
            "situational_patterns": []
        }
    
    # Модуль 12: Управленческие рекомендации (синтез)
    management_data = prepare_management_data(intermediate_result)
    
    # ЭТАП 6: Формирование итоговой выжимки
    logger.info("Этап 6: Формирование итоговой выжимки...")
    
    result = {
        "metadata": {
            "participant_name": participant_name,
            "exercises_count": len(exercises),
            "total_words": len(all_text.split()),
            "processing_timestamp": "2025-01-09"
        },
        "assessor_discussion": assessor_discussion,  # НОВОЕ!
        "disfluency_profile": disfluency_profile,
        "critical_incidents": critical_incidents,
        "mbti_markers": mbti_result,
        "pavlov_markers": pavlov_result,
        "ocean_markers": ocean_result,
        "radicals_markers": radicals_result,
        "dark_tetrad_markers": dark_tetrad_result,
        "potential_markers": potential_result,
        "stress_profile": stress_result,
        "shadow_sides": shadow_result,
        "prism_layers": prism_result,
        "competencies_analysis": competencies_result,
        "red_thread": red_thread_result,
        "management_prep": management_data,
        "psycholinguistics": psycholing_result,
        "base_metrics": {
            "total_words": len(all_text.split()),
            "unique_words": len(set(lemmatize_text(clean))),
            "avg_sentence_length": _calculate_avg_sentence_length(clean)
        },
        "participant_summary": _create_participant_summary(clean, critical_incidents),  # ✅ Сжатая выжимка (для /send)
        "participant_full_texts": []  # ✅ ПОЛНЫЕ тексты участника (для /export) - заполняется извне
    }
    
    logger.info("✅ Обработка завершена - ВСЕ 12 модулей проанализированы")
    return result


# ============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ============================================================================

def _calculate_avg_sentence_length(text: str) -> float:
    """Средняя длина предложения"""
    sentences = re.split(r'[.!?]+', text)
    sentences = [s for s in sentences if s.strip()]
    if not sentences:
        return 0.0
    
    total_words = sum(len(s.split()) for s in sentences)
    return round(total_words / len(sentences), 1)


def _create_participant_summary(text: str, critical_incidents: List[Dict]) -> str:
    """
    Создает сжатую выжимку речи участника (600-800 слов)
    Берет начало, середину, конец + критические инциденты
    """
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    
    if len(sentences) < 10:
        return text[:5000]  # Если мало текста, возвращаем как есть
    
    # Берем первые 10 предложений (начало)
    intro = ' '.join(sentences[:10])
    
    # Берем 10 из середины
    mid_start = len(sentences) // 3
    middle = ' '.join(sentences[mid_start:mid_start+10])
    
    # Берем последние 5 (финал)
    outro = ' '.join(sentences[-5:])
    
    # Добавляем тексты критических инцидентов
    incidents_text = ' '.join([inc.get('text', '') for inc in critical_incidents[:5]])
    
    summary = f"{intro} [...] {incidents_text} [...] {middle} [...] {outro}"
    
    # Ограничиваем 800 словами
    words = summary.split()
    if len(words) > 800:
        summary = ' '.join(words[:800]) + '...'
    
    return summary

def _extract_top_quotes(text: str, max_quotes: int = 5) -> List[str]:
    """Извлечение самых информативных цитат"""
    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if 5 <= len(s.split()) <= 30]
    
    # Скорим предложения по информативности
    scored = []
    for sent in sentences:
        score = 0
        sent_lower = sent.lower()
        
        # Высокий приоритет: самохарактеристики
        if re.search(r'\bя\b.*(думаю|считаю|чувствую|стремлюсь|предпочитаю)', sent_lower):
            score += 10
        
        # Средний приоритет: действия и решения
        if re.search(r'\b(решил|сделал|организовал|создал|достиг)', sent_lower):
            score += 5
        
        # Низкий приоритет: общие утверждения
        if re.search(r'\b(важно|ценю|стараюсь|обычно)', sent_lower):
            score += 3
        
        scored.append((score, sent))
    
    # Сортируем и берем топ
    scored.sort(reverse=True, key=lambda x: x[0])
    return [s[1] for s in scored[:max_quotes]]


# ============================================================================
# ТЕСТИРОВАНИЕ
# ============================================================================

if __name__ == "__main__":
    test_text = """
    Я считаю важным тщательно анализировать ситуацию перед принятием решения. 
    Обычно я стараюсь рассмотреть все детали и возможные последствия. В команде 
    я предпочитаю роль организатора, помогаю структурировать работу. Ценю порядок 
    и систематичность. Когда возникают сложности, я спокойно ищу решение.
    
    Один раз в проекте X мне пришлось организовать работу команды из 5 человек.
    Я разработал план, распределил задачи и контролировал выполнение. В итоге
    проект был завершен в срок.
    """
    
    result = process_exercises([test_text], "Тестовый участник")
    print(json.dumps(result, ensure_ascii=False, indent=2))

